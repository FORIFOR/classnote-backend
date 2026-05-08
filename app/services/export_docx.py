"""DOCX renderer for session export.

Generates a Word document that preserves every field of the summary JSON.
Matches the PDF export layout: hero callout, participants, highlights with
category chips, section blocks (text / bullets / grouped_bullets / table /
cards), keywords, optional transcript appendix.
"""
import io
import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

# ───── Colors ─────
C_PRIMARY   = RGBColor(0x1F, 0x4F, 0xD8)
C_PRIMARY_BG = "EEF3FF"
C_HEAD      = RGBColor(0x1A, 0x1A, 0x2E)
C_TEXT      = RGBColor(0x11, 0x11, 0x11)
C_MUTED     = RGBColor(0x88, 0x88, 0x99)
C_CAPTION   = RGBColor(0x55, 0x55, 0x66)
C_HAIR      = "E4E4EC"
C_CARD_BG   = "F7F8FB"
C_TH_BG     = "EDEEF5"

EMPHASIS = {
    "primary": ("1F4FD8", "EEF3FF"),
    "success": ("16A34A", "ECFDF3"),
    "warning": ("EA580C", "FFF4ED"),
    "neutral": ("6B7280", "F3F4F6"),
}


def render_docx(
    *,
    title: str,
    date_text: Optional[str] = None,
    duration_text: Optional[str] = None,
    mode: str = "meeting",
    highlights: Optional[List[Any]] = None,
    overview: Optional[str] = None,
    keywords: Optional[List[str]] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    transcript: Optional[str] = None,
    bottom_line: Optional[str] = None,
    why_it_matters: Optional[str] = None,
    outcome_status: Optional[str] = None,
    participants: Optional[List[Dict[str, str]]] = None,
    **_kwargs,
) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Kaku Gothic Pro"
    style.font.size = Pt(10.5)
    _set_east_asia_font(style, "Hiragino Kaku Gothic Pro")

    # ── Title ──
    ttl = doc.add_paragraph()
    run = ttl.add_run(title or "Untitled")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = C_HEAD

    # ── Meta line ──
    meta_parts = [p for p in [date_text, duration_text, _mode_label(mode)] if p]
    if meta_parts:
        mp = doc.add_paragraph()
        mr = mp.add_run(" · ".join(meta_parts))
        mr.font.size = Pt(9)
        mr.font.color.rgb = C_CAPTION
        mp.paragraph_format.space_after = Pt(8)

    # ── Hero callout ──
    if bottom_line or why_it_matters:
        _hero_callout(doc, bottom_line, why_it_matters, outcome_status)

    # ── Participants ──
    if participants:
        _section_header(doc, "参加者", "1F4FD8")
        _participants_block(doc, participants)

    # ── Highlights ──
    rich_hls = _normalize_highlights(highlights)
    if rich_hls:
        _section_header(doc, "要点", "1F4FD8")
        for h in rich_hls:
            _highlight_row(doc, h)

    # ── Overview ──
    if overview:
        _section_header(doc, "概要", "1F4FD8")
        for para in str(overview).split("\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(4)

    # ── Sections ──
    for sec in (sections or []):
        _render_section(doc, sec)

    # ── Keywords ──
    if keywords:
        _section_header(doc, "キーワード", "1F4FD8")
        _keywords_block(doc, keywords)

    # ── Transcript appendix ──
    if transcript:
        doc.add_page_break()
        _section_header(doc, "文字起こし", "1F4FD8")
        for line in str(transcript).split("\n"):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                pr = p.runs[0]
                pr.font.size = Pt(9)

    # ── Footer ──
    footer = doc.sections[-1].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fr = fp.add_run("Created by DeepNote")
    fr.font.size = Pt(8)
    fr.font.color.rgb = C_MUTED
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────
# Section dispatcher
# ─────────────────────────────────────────────────────
def _render_section(doc: Document, sec: Dict[str, Any]) -> None:
    sec_type = sec.get("type", "text")
    sec_title = sec.get("title", "")
    emphasis = sec.get("emphasis", "primary")
    accent, _bg = EMPHASIS.get(emphasis, EMPHASIS["primary"])

    if sec_title:
        _section_header(doc, sec_title, accent)

    if sec_type == "text":
        body = sec.get("body", "")
        for para in str(body).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    elif sec_type == "bullets":
        for item in sec.get("items", []):
            p = doc.add_paragraph(str(item), style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

    elif sec_type == "grouped_bullets":
        for g in sec.get("groups", []):
            label = g.get("label")
            if label:
                lp = doc.add_paragraph()
                lr = lp.add_run(label)
                lr.bold = True
                lr.font.size = Pt(10)
                lr.font.color.rgb = RGBColor.from_string(accent)
                lp.paragraph_format.space_before = Pt(3)
                lp.paragraph_format.space_after = Pt(1)
            for item in g.get("items", []):
                p = doc.add_paragraph(str(item), style="List Bullet")
                p.paragraph_format.space_after = Pt(2)

    elif sec_type == "table":
        columns = sec.get("columns", [])
        rows = sec.get("rows", [])
        if columns and rows:
            _rich_table(doc, columns, rows, accent)

    elif sec_type == "cards":
        for card in sec.get("cards", []):
            _card_block(doc, card, accent)


# ─────────────────────────────────────────────────────
# Components
# ─────────────────────────────────────────────────────
def _section_header(doc: Document, label: str, accent_hex: str) -> None:
    """Plain bold mid-section heading with a thin hairline underline.

    2026-05-08 design refresh: previously this rendered a 1x1 colored
    bar + heading combo via a table. The colored stripe felt visually
    heavy ("中項目の青いところがいらない"); the new layout is just a
    bold paragraph with a 0.5pt grey bottom border so sections still
    have a clear separator without competing with the body text.

    ``accent_hex`` is intentionally ignored to keep the visual language
    consistent (only the hero callout retains the brand accent).
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_HEAD
    # Hairline grey underline at the paragraph level
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")          # 0.5pt
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D6D6DC")  # very light grey
    pBdr.append(bottom)
    pPr.append(pBdr)


def _hero_callout(
    doc: Document,
    bottom_line: Optional[str],
    why_it_matters: Optional[str],
    outcome_status: Optional[str],
) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17.0)
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, C_PRIMARY_BG)
    _set_cell_margins(cell, top=200, left=240, bottom=200, right=240)
    _set_cell_left_border(cell, "1F4FD8", sz=24)

    cell.text = ""
    if outcome_status:
        p = cell.paragraphs[0]
        r = p.add_run(_outcome_label(outcome_status))
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_PRIMARY
        p.paragraph_format.space_after = Pt(3)
        bl_p = cell.add_paragraph()
    else:
        bl_p = cell.paragraphs[0]

    if bottom_line:
        r = bl_p.add_run(bottom_line)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = C_HEAD
        bl_p.paragraph_format.space_after = Pt(2)

    if why_it_matters:
        wp = cell.add_paragraph()
        wr = wp.add_run(why_it_matters)
        wr.font.size = Pt(9)
        wr.font.color.rgb = C_CAPTION

    _remove_table_borders(tbl, keep_left=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _participants_block(doc: Document, participants: List[Dict[str, str]]) -> None:
    cols = 3
    rows_count = (len(participants) + cols - 1) // cols
    tbl = doc.add_table(rows=rows_count, cols=cols)
    tbl.autofit = False
    for col in tbl.columns:
        col.width = Cm(5.6)
    for idx, p in enumerate(participants):
        r, c = idx // cols, idx % cols
        cell = tbl.rows[r].cells[c]
        _shade_cell(cell, C_CARD_BG)
        cell.text = ""
        para = cell.paragraphs[0]
        name_run = para.add_run(p.get("name") or "(不明)")
        name_run.bold = True
        name_run.font.size = Pt(9.5)
        if p.get("role"):
            role_run = para.add_run(f"  {p['role']}")
            role_run.font.size = Pt(8.5)
            role_run.font.color.rgb = C_MUTED
        _set_cell_margins(cell, top=80, left=120, bottom=80, right=120)
    _apply_hairline_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _highlight_row(doc: Document, h: Dict[str, Any]) -> None:
    p = doc.add_paragraph(style="List Bullet")
    category = h.get("category")
    if category:
        cr = p.add_run(f"[{_category_label(category)}] ")
        cr.bold = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = C_PRIMARY
    tr = p.add_run(h.get("text") or "")
    tr.font.size = Pt(10)
    tr.font.color.rgb = C_TEXT
    if h.get("needConfirm"):
        nr = p.add_run("  [要確認]")
        nr.font.size = Pt(8.5)
        nr.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
    p.paragraph_format.space_after = Pt(3)


def _rich_table(doc: Document, columns: List[str], rows: List[list], accent_hex: str) -> None:
    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i, col in enumerate(columns):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, "EDEEF5")
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(col)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_HEAD
        _set_cell_margins(cell, top=60, left=100, bottom=60, right=100)

    # Body
    for row_data in rows:
        row = tbl.add_row()
        padded = list(row_data) + [""] * (len(columns) - len(row_data))
        for i, cell_text in enumerate(padded[: len(columns)]):
            cell = row.cells[i]
            cell.text = ""
            for j, line in enumerate(str(cell_text or "").split("\n")):
                if j == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                r = p.add_run(line)
                r.font.size = Pt(9)
                r.font.color.rgb = C_TEXT
            _set_cell_margins(cell, top=60, left=100, bottom=60, right=100)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _card_block(doc: Document, card: Dict[str, Any], accent_hex: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17.0)
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, C_CARD_BG)
    _set_cell_margins(cell, top=140, left=180, bottom=140, right=180)
    _set_cell_left_border(cell, accent_hex, sz=16)

    cell.text = ""
    title = card.get("title", "")
    first_para = cell.paragraphs[0]
    if title:
        tr = first_para.add_run(title)
        tr.bold = True
        tr.font.size = Pt(10.5)
        tr.font.color.rgb = C_HEAD
        first_para.paragraph_format.space_after = Pt(2)
    written = bool(title)
    for key, value in card.get("fields", []):
        if not value:
            continue
        p = cell.add_paragraph() if written or title else first_para
        if key:
            kr = p.add_run(f"{key}  ")
            kr.font.size = Pt(8.5)
            kr.font.color.rgb = C_MUTED
        vr = p.add_run(str(value))
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = C_TEXT
        p.paragraph_format.space_after = Pt(1)
        written = True
    if not written:
        first_para.add_run("—").font.size = Pt(9)

    _apply_hairline_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def _keywords_block(doc: Document, keywords: List[str]) -> None:
    p = doc.add_paragraph()
    for i, k in enumerate(keywords):
        if i > 0:
            sep = p.add_run("   ")
            sep.font.size = Pt(9)
        r = p.add_run(f" {k} ")
        r.font.size = Pt(9)
        r.font.color.rgb = C_PRIMARY
        r.bold = True
    p.paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────────────
def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_margins(cell, *, top=0, left=0, bottom=0, right=0) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_cell_left_border(cell, color_hex: str, sz: int = 16) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:color"), color_hex)
    borders.append(left)
    for side in ("top", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tcPr.append(borders)


def _remove_table_borders(tbl, keep_left: bool = False) -> None:
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                if keep_left and side == "left":
                    continue
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:val"), "nil")
                borders.append(node)
            tcPr.append(borders)


def _apply_hairline_borders(tbl) -> None:
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:val"), "single")
                node.set(qn("w:sz"), "4")
                node.set(qn("w:color"), C_HAIR)
                borders.append(node)
            tcPr.append(borders)


def _set_east_asia_font(style, font_name: str) -> None:
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# ─────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────
def _normalize_highlights(highlights) -> List[Dict[str, Any]]:
    if not highlights:
        return []
    out = []
    for h in highlights:
        if isinstance(h, str):
            out.append({"text": h, "category": None, "needConfirm": False})
        elif isinstance(h, dict):
            out.append({
                "text": h.get("text") or h.get("title") or "",
                "category": h.get("category"),
                "needConfirm": bool(h.get("needConfirm")),
            })
    return [x for x in out if x.get("text")]


def _mode_label(mode: str) -> str:
    return {"lecture": "講義", "meeting": "会議", "translate": "翻訳"}.get(mode, mode)


def _category_label(cat: str) -> str:
    return {
        "decision": "決定",
        "todo": "TODO",
        "concern": "懸念",
        "insight": "気づき",
        "fact": "事実",
        "info": "情報",
        "risk": "リスク",
    }.get(cat, cat)


def _outcome_label(status: str) -> str:
    return {
        "action_agreed": "✓ 合意・次アクション確定",
        "tentative_alignment": "△ 暫定合意",
        "discussion_only": "・ 議論のみ",
        "unresolved": "! 未解決",
    }.get(status, status)
